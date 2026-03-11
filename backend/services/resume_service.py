import os
import tempfile
from typing import Optional
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from PyPDF2 import PdfReader
from docx import Document
from database import get_db
from models.user import User
from models.resume import Resume
from ai_agents.resume_analyzer import ResumeAnalyzer

class ResumeService:
    def __init__(self):
        self.analyzer = ResumeAnalyzer()
        self.upload_dir = "uploads/resumes"
        os.makedirs(self.upload_dir, exist_ok=True)
    
    async def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with multiple fallback methods."""
        text = ""
        
        # Method 1: Try PyPDF2 first (fastest)
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                            text += page_text + "\n"
                    except Exception:
                        continue
                
                if text.strip() and len(text.strip()) > 100:
                    return text.strip()
        except Exception:
            pass
        
        # Method 2: Try pdfplumber (better for complex PDFs)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                            text += page_text + "\n"
                    except Exception:
                        continue
                
                if text.strip() and len(text.strip()) > 100:
                    return text.strip()
        except ImportError:
            pass  # pdfplumber not installed
        except Exception:
            pass
        
        # Method 3: Try OCR for scanned PDFs
        try:
            ocr_text = await self._extract_text_with_ocr(file_path)
            if ocr_text and len(ocr_text.strip()) > 50:
                return ocr_text.strip()
        except Exception:
            pass
        
        # If all methods failed
        if text.strip() and len(text.strip()) > 20:
            return text.strip()
        
        raise HTTPException(
            status_code=400,
            detail="Could not extract sufficient text from PDF. The file might be: 1) Scanned/image-based (OCR not available), 2) Password protected, 3) Corrupted, or 4) Empty. Please try: 1) Converting to a text-based PDF, 2) Using a DOCX file instead, or 3) Ensuring the PDF contains selectable text."
        )
    
    async def _extract_text_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR (for scanned/image-based PDFs)."""
        try:
            # Try to import OCR libraries
            try:
                from pdf2image import convert_from_path
                import pytesseract
                from PIL import Image
            except ImportError:
                # OCR libraries not available
                return ""
            
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300, first_page=1, last_page=5)  # Limit to first 5 pages
            
            text = ""
            for i, image in enumerate(images):
                try:
                    # Extract text from image using OCR
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    if page_text:
                        page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                        text += page_text + "\n"
                except Exception:
                    continue
            
            return text.strip()
            
        except Exception as e:
            # OCR failed, return empty string
            return ""
    
    async def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                try:
                    # Clean up text by handling encoding issues
                    para_text = paragraph.text
                    if para_text:
                        para_text = para_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                        text += para_text + "\n"
                except Exception:
                    continue
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        try:
                            cell_text = cell.text
                            if cell_text:
                                cell_text = cell_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                                text += cell_text + " "
                        except Exception:
                            continue
                text += "\n"
            
            if not text.strip():
                raise HTTPException(status_code=400, detail="Could not extract text from DOCX. The file might be empty or corrupted.")
            
            return text.strip()
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")
    
    async def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text based on file type."""
        if file_extension.lower() == '.pdf':
            return await self.extract_text_from_pdf(file_path)
        elif file_extension.lower() in ['.docx', '.doc']:
            return await self.extract_text_from_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX files.")
    
    async def save_uploaded_file(self, file: UploadFile, user_id: int) -> tuple[str, str]:
        """Save uploaded file and return file path and filename."""
        # Generate unique filename
        file_extension = os.path.splitext(file.filename)[1]
        safe_filename = f"user_{user_id}_{file.filename}"
        file_path = os.path.join(self.upload_dir, safe_filename)
        
        # Save file
        try:
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            return file_path, safe_filename
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    
    async def process_resume(self, file: UploadFile, user_id: int, db: Session) -> Resume:
        """Process uploaded resume and save to database."""
        # Validate file type
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ['.pdf', '.docx', '.doc']:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        # Validate file size (max 10MB)
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > 10 * 1024 * 1024:  # 10MB
            raise HTTPException(status_code=400, detail="File size must be less than 10MB")
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        
        # Save file
        file_path, filename = await self.save_uploaded_file(file, user_id)
        
        try:
            # Extract text
            extracted_text = await self.extract_text(file_path, file_extension)
            
            # Validate extracted text
            if not extracted_text or len(extracted_text.strip()) < 50:
                raise HTTPException(
                    status_code=400, 
                    detail="Could not extract sufficient text from the file. Please ensure the PDF is not scanned or image-based."
                )
            
            # Analyze resume
            analysis_result = self.analyzer.analyze_resume(extracted_text)
            
            # Save to database
            db_resume = Resume(
                user_id=user_id,
                filename=filename,
                file_path=file_path,
                extracted_text=extracted_text,
                analysis_result=analysis_result.dict(),
                score=analysis_result.score
            )
            
            db.add(db_resume)
            db.commit()
            db.refresh(db_resume)
            
            return db_resume
            
        except HTTPException:
            # Clean up file if processing failed
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            raise
        except Exception as e:
            # Clean up file if processing failed
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")
    
    def get_user_resumes(self, user_id: int, db: Session) -> list[Resume]:
        """Get all resumes for a user."""
        return db.query(Resume).filter(Resume.user_id == user_id).all()
    
    def get_resume_by_id(self, resume_id: int, user_id: int, db: Session) -> Optional[Resume]:
        """Get a specific resume by ID."""
        return db.query(Resume).filter(
            Resume.id == resume_id,
            Resume.user_id == user_id
        ).first()
